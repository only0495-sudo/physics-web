from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
ASSET_DIR = ROOT / "ch4_worksheet_assets_tmp"
OUTPUT = ROOT / "高一物理_CH4_學生學習單.docx"

ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_CJK = "Microsoft JhengHei"
FONT_LATIN = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F5F8FC"
LIGHT_GRAY = "F2F4F7"
BORDER = "B9C7D8"
WHITE = "FFFFFF"
ACCENT = "D97706"
GREEN = "2E7D65"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 80
CELL_START_END = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None,
                 name: str = FONT_CJK):
    run.font.name = name
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_LATIN if name == FONT_CJK else name)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN if name == FONT_CJK else name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_style_font(style, name=FONT_CJK, size=11, color=INK, bold=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=CELL_TOP_BOTTOM, start=CELL_START_END,
                     bottom=CELL_TOP_BOTTOM, end=CELL_START_END):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names += ["insideH", "insideV"]
    for edge in names:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_shading(paragraph, fill=LIGHTER_BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, edge="bottom", color=BORDER, size=6, space=2):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color=INK, size=10.2,
                  align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        set_cell_shading(cell, fill)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(document: Document, ordered=False):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(rfonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    keep.set(qn("w:val"), "1" if value else "0")


def set_keep_together(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepLines"))
    if keep is None:
        keep = OxmlElement("w:keepLines")
        p_pr.append(keep)
    keep.set(qn("w:val"), "1" if value else "0")


def add_field(paragraph, field_code: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_run_font(run, size=9, color=MUTED)


def add_body(doc, text="", bold=False, color=INK, size=None, after=6,
             before=0, align=WD_ALIGN_PARAGRAPH.LEFT, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size or 11, bold=bold, color=color)
    if keep:
        set_keep_together(p)
    return p


def add_fill(doc, prompt, lines=1, compact=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(prompt)
    set_run_font(r, size=10.5 if compact else 11, color=INK)
    for _ in range(lines):
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Inches(0.18)
        line.paragraph_format.right_indent = Inches(0.12)
        line.paragraph_format.space_before = Pt(3)
        line.paragraph_format.space_after = Pt(5)
        line.paragraph_format.line_spacing = 1.0
        line.add_run("\u00A0")
        set_paragraph_border(line, "bottom", color="AAB7C8", size=5, space=2)
    return p


def add_callout(doc, label, text, fill=LIGHTER_BLUE, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, "left", color=color, size=18, space=6)
    r1 = p.add_run(f"{label}　")
    set_run_font(r1, size=10.5, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_picture(doc, path, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    shape = r.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", path.stem.replace("_", " "))
    return p


def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    apply_numbering(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a, size=10.7, bold=True, color=INK)
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b, size=10.7, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.7, color=INK)
    return p


def add_question_group(doc, title, questions):
    add_section_heading(doc, title, 3)
    num_id = add_numbering_definition(doc, ordered=True)
    for item in questions:
        if isinstance(item, tuple):
            text, lines = item
        else:
            text, lines = item, 1
        p = add_list_item(doc, text, num_id)
        p.paragraph_format.space_after = Pt(3)
        for _ in range(lines):
            answer = doc.add_paragraph()
            answer.paragraph_format.left_indent = Inches(0.38)
            answer.paragraph_format.right_indent = Inches(0.1)
            answer.paragraph_format.space_before = Pt(3)
            answer.paragraph_format.space_after = Pt(5)
            answer.add_run("\u00A0")
            set_paragraph_border(answer, "bottom", color="B8C2CF", size=5, space=2)


def header_row(table, labels, fill=LIGHT_BLUE):
    row = table.rows[0]
    for cell, label in zip(row.cells, labels):
        set_cell_text(cell, label, bold=True, color=DARK_BLUE, size=9.7,
                      align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
    set_repeat_table_header(row)


def add_table(doc, headers, rows, widths, font_size=9.5, fills=None,
              alignments=None, header_fill=LIGHT_BLUE, indent=TABLE_INDENT_DXA):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent)
    set_table_borders(table)
    header_row(table, headers, header_fill)
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        for c_idx, value in enumerate(values):
            align = (
                alignments[c_idx]
                if alignments and c_idx < len(alignments)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            fill = fills[r_idx][c_idx] if fills else None
            set_cell_text(row.cells[c_idx], str(value), size=font_size,
                          align=align, fill=fill)
    set_table_geometry(table, widths, indent)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def pil_font(size, bold=False):
    path = Path("C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc")
    return ImageFont.truetype(str(path), size=size)


def draw_arrow(draw, start, end, fill="#1F4D78", width=6, head=18):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    left = (
        end[0] - head * math.cos(angle - math.pi / 6),
        end[1] - head * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - head * math.cos(angle + math.pi / 6),
        end[1] - head * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=fill)


def draw_centered(draw, xy, text, font, fill="#243447"):
    box = draw.textbbox((0, 0), text, font=font)
    x = xy[0] - (box[2] - box[0]) / 2
    y = xy[1] - (box[3] - box[1]) / 2
    draw.text((x, y), text, font=font, fill=fill)


def diagram_current_wire(path):
    img = Image.new("RGB", (1600, 560), "white")
    d = ImageDraw.Draw(img)
    title = pil_font(36, True)
    label = pil_font(28)
    small = pil_font(23)
    d.rounded_rectangle((15, 15, 1585, 545), radius=28, outline="#B9C7D8", width=4)
    for panel, center, symbol, clockwise in [
        ((45, 80, 770, 510), (405, 295), "⊙", False),
        ((830, 80, 1555, 510), (1190, 295), "⊗", True),
    ]:
        d.rounded_rectangle(panel, radius=22, fill="#F5F8FC", outline="#D5DFEA", width=3)
        draw_centered(d, ((panel[0] + panel[2]) / 2, 120),
                      "直導線截面", title, "#1F4D78")
        d.ellipse((center[0] - 58, center[1] - 58, center[0] + 58, center[1] + 58),
                  fill="white", outline="#2E74B5", width=6)
        draw_centered(d, center, symbol, pil_font(64, True), "#D97706")
        radius = 145
        bbox = (center[0] - radius, center[1] - radius, center[0] + radius, center[1] + radius)
        d.arc(bbox, start=35, end=325, fill="#2E7D65", width=8)
        theta = math.radians(325 if not clockwise else 35)
        if clockwise:
            start = (
                center[0] + radius * math.cos(theta + 0.15),
                center[1] + radius * math.sin(theta + 0.15),
            )
            end = (
                center[0] + radius * math.cos(theta),
                center[1] + radius * math.sin(theta),
            )
        else:
            start = (
                center[0] + radius * math.cos(theta - 0.15),
                center[1] + radius * math.sin(theta - 0.15),
            )
            end = (
                center[0] + radius * math.cos(theta),
                center[1] + radius * math.sin(theta),
            )
        draw_arrow(d, start, end, fill="#2E7D65", width=8, head=24)
        draw_centered(d, ((panel[0] + panel[2]) / 2, 465),
                      "磁場方向：________________", label)
        draw_centered(d, ((panel[0] + panel[2]) / 2, 410),
                      "電流方向：________________", small, "#667085")
    img.save(path, dpi=(180, 180))


def diagram_lenz(path):
    img = Image.new("RGB", (1600, 650), "white")
    d = ImageDraw.Draw(img)
    title = pil_font(33, True)
    label = pil_font(24)
    small = pil_font(21)
    d.rounded_rectangle((15, 15, 1585, 635), radius=28, outline="#B9C7D8", width=4)
    panels = [
        (45, 70, 770, 595, "N 極靠近線圈", True),
        (830, 70, 1555, 595, "N 極遠離線圈", False),
    ]
    for x1, y1, x2, y2, heading, approaching in panels:
        d.rounded_rectangle((x1, y1, x2, y2), radius=22, fill="#F5F8FC",
                            outline="#D5DFEA", width=3)
        draw_centered(d, ((x1 + x2) / 2, 112), heading, title, "#1F4D78")
        magnet_x = x1 + 115 if approaching else x1 + 235
        magnet_y = 235
        d.rectangle((magnet_x, magnet_y, magnet_x + 220, magnet_y + 92),
                    fill="#FFFFFF", outline="#374151", width=4)
        d.rectangle((magnet_x + 110, magnet_y, magnet_x + 220, magnet_y + 92),
                    fill="#DDE7F3")
        draw_centered(d, (magnet_x + 55, magnet_y + 46), "S", pil_font(35, True), "#374151")
        draw_centered(d, (magnet_x + 165, magnet_y + 46), "N", pil_font(35, True), "#B42318")
        coil_x = x1 + 515
        for off in range(-45, 46, 15):
            d.ellipse((coil_x + off, 185, coil_x + off + 68, 375),
                      outline="#2E74B5", width=5)
        if approaching:
            draw_arrow(d, (magnet_x + 235, 281), (coil_x - 35, 281),
                       fill="#D97706", width=7, head=24)
        else:
            draw_arrow(d, (magnet_x - 15, 281), (magnet_x - 115, 281),
                       fill="#D97706", width=7, head=24)
        draw_centered(d, ((x1 + x2) / 2, 430),
                      "線圈靠磁鐵側：____ 極", label)
        draw_centered(d, ((x1 + x2) / 2, 480),
                      "感應磁場方向：____________", label)
        draw_centered(d, ((x1 + x2) / 2, 530),
                      "由磁鐵側看電流：____________", small, "#667085")
    img.save(path, dpi=(180, 180))


def diagram_wave(path):
    img = Image.new("RGB", (1600, 560), "white")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((15, 15, 1585, 545), radius=28, outline="#B9C7D8", width=4)
    axis_y = 310
    d.line((90, axis_y, 1510, axis_y), fill="#9AA7B6", width=3)
    points = []
    for x in range(100, 1510, 4):
        y = axis_y - 130 * math.sin((x - 100) * 2 * math.pi / 520)
        points.append((x, y))
    d.line(points, fill="#2E74B5", width=8)
    peak1 = (230, 180)
    peak2 = (750, 180)
    draw_arrow(d, (230, 135), (750, 135), fill="#D97706", width=5, head=18)
    draw_arrow(d, (750, 135), (230, 135), fill="#D97706", width=5, head=18)
    draw_centered(d, (490, 92), "① ____________________", pil_font(28, True), "#D97706")
    draw_arrow(d, (955, axis_y), (955, 180), fill="#2E7D65", width=5, head=18)
    draw_arrow(d, (955, 180), (955, axis_y), fill="#2E7D65", width=5, head=18)
    d.text((985, 220), "② ____________________", font=pil_font(28, True), fill="#2E7D65")
    d.text((100, 430), "平衡位置", font=pil_font(25), fill="#667085")
    d.text((1160, 430), "波向右傳播 →", font=pil_font(29, True), fill="#1F4D78")
    img.save(path, dpi=(180, 180))


def diagram_optics(path):
    img = Image.new("RGB", (1600, 650), "white")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((15, 15, 1585, 635), radius=28, outline="#B9C7D8", width=4)
    title = pil_font(33, True)
    label = pil_font(24)
    for x1, x2, heading in [(45, 770, "反射"), (830, 1555, "折射")]:
        d.rounded_rectangle((x1, 70, x2, 595), radius=22, fill="#F5F8FC",
                            outline="#D5DFEA", width=3)
        draw_centered(d, ((x1 + x2) / 2, 112), heading, title, "#1F4D78")
        center = ((x1 + x2) // 2, 330)
        d.line((x1 + 90, center[1], x2 - 90, center[1]), fill="#374151", width=5)
        for y in range(150, 525, 22):
            d.line((center[0], y, center[0], y + 11), fill="#7B8794", width=3)
        draw_arrow(d, (x1 + 145, 160), center, fill="#D97706", width=7, head=23)
        if heading == "反射":
            draw_arrow(d, center, (x2 - 145, 160), fill="#2E74B5", width=7, head=23)
            d.text((x1 + 115, 500), "入射角：____　反射角：____",
                   font=label, fill="#243447")
        else:
            draw_arrow(d, center, (x2 - 210, 535), fill="#2E74B5", width=7, head=23)
            d.text((x1 + 115, 500), "介質 1：n₁ = ____　介質 2：n₂ = ____",
                   font=pil_font(21), fill="#243447")
        d.text((center[0] + 15, 150), "法線", font=label, fill="#667085")
    img.save(path, dpi=(180, 180))


def diagram_diffraction_doppler(path):
    img = Image.new("RGB", (1600, 650), "white")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((15, 15, 1585, 635), radius=28, outline="#B9C7D8", width=4)
    title = pil_font(31, True)
    label = pil_font(23)
    panels = [(45, 70, 770, 595), (830, 70, 1555, 595)]
    for p in panels:
        d.rounded_rectangle(p, radius=22, fill="#F5F8FC", outline="#D5DFEA", width=3)

    draw_centered(d, (408, 112), "狹縫繞射", title, "#1F4D78")
    for x in (125, 195, 265, 335):
        d.line((x, 170, x, 500), fill="#2E74B5", width=5)
    d.line((430, 155, 430, 285), fill="#374151", width=14)
    d.line((430, 385, 430, 515), fill="#374151", width=14)
    center = (430, 335)
    for radius in (85, 150, 220):
        d.arc((center[0] - radius, center[1] - radius,
               center[0] + radius, center[1] + radius),
              start=-70, end=70, fill="#2E7D65", width=5)
    d.text((110, 535), "當 λ 與狹縫寬度接近時：____________",
           font=label, fill="#243447")

    draw_centered(d, (1192, 112), "都卜勒波前", title, "#1F4D78")
    source = (1190, 325)
    d.ellipse((source[0] - 25, source[1] - 25, source[0] + 25, source[1] + 25),
              fill="#D97706", outline="#9A5A00", width=3)
    draw_arrow(d, (1180, 215), (1370, 215), fill="#D97706", width=7, head=24)
    d.text((1210, 165), "波源移動方向", font=label, fill="#D97706")
    for radius in (90, 165, 245, 325):
        d.arc((source[0] - radius * 1.4, source[1] - radius,
               source[0] + radius * 0.55, source[1] + radius),
              start=105, end=255, fill="#2E74B5", width=4)
    for radius in (55, 100, 145, 190):
        d.arc((source[0] - radius * 0.35, source[1] - radius,
               source[0] + radius, source[1] + radius),
              start=-75, end=75, fill="#2E7D65", width=4)
    d.text((875, 535), "左側 λ：____　右側 λ：____　右側 f：____",
           font=label, fill="#243447")
    img.save(path, dpi=(180, 180))


def build_diagrams():
    paths = {
        "wire": ASSET_DIR / "current_wire.png",
        "lenz": ASSET_DIR / "lenz_law.png",
        "wave": ASSET_DIR / "wave_quantities.png",
        "optics": ASSET_DIR / "reflection_refraction.png",
        "diffraction_doppler": ASSET_DIR / "diffraction_doppler.png",
    }
    diagram_current_wire(paths["wire"])
    diagram_lenz(paths["lenz"])
    diagram_wave(paths["wave"])
    diagram_optics(paths["optics"])
    diagram_diffraction_doppler(paths["diffraction_doppler"])
    return paths


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        header = sec.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        r = hp.add_run("高一物理｜CH.4 電與磁的統一｜學生學習單")
        set_run_font(r, size=8.5, color=MUTED, bold=True)

        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run("CH.4　｜　")
        set_run_font(r, size=8.5, color=MUTED)
        add_field(fp, "PAGE")


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("STUDENT WORKSHEET  ·  PHYSICS 1")
    set_run_font(r, size=10, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("高一物理：CH.4")
    set_run_font(r, size=29, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("電與磁的統一")
    set_run_font(r, size=22, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("從動電生磁、動磁生電，到光與波的共同語言")
    set_run_font(r, size=12.5, color=MUTED)

    info = doc.add_table(rows=1, cols=4)
    set_table_geometry(info, [1150, 2750, 1150, 4310], indent_dxa=0)
    set_table_borders(info, color=BORDER, size=6)
    labels = ["班級", "____________", "姓名／座號", "____________________________"]
    for i, label in enumerate(labels):
        set_cell_text(
            info.rows[0].cells[i],
            label,
            bold=i % 2 == 0,
            color=DARK_BLUE if i % 2 == 0 else INK,
            size=10.5,
            align=WD_ALIGN_PARAGRAPH.CENTER if i % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            fill=LIGHT_BLUE if i % 2 == 0 else WHITE,
        )
    set_table_geometry(info, [1150, 2750, 1150, 4310], indent_dxa=0)
    add_body(doc, "", after=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("章節地圖")
    set_run_font(r, size=12, bold=True, color=DARK_BLUE)

    roadmap = doc.add_table(rows=2, cols=5)
    widths = [1872] * 5
    set_table_geometry(roadmap, widths, indent_dxa=0)
    set_table_borders(roadmap, color="C8D5E4", size=5)
    units = [
        ("01", "電流磁效應"),
        ("02", "電磁感應"),
        ("03", "電磁波"),
        ("04", "波動與光學"),
        ("05", "都卜勒效應"),
    ]
    for i, (num, title) in enumerate(units):
        set_cell_text(roadmap.rows[0].cells[i], num, bold=True, color=WHITE,
                      size=11, align=WD_ALIGN_PARAGRAPH.CENTER, fill=BLUE)
        set_cell_text(roadmap.rows[1].cells[i], title, bold=True, color=DARK_BLUE,
                      size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, fill=LIGHTER_BLUE)
    set_table_geometry(roadmap, widths, indent_dxa=0)
    add_body(doc, "", after=2)

    add_callout(
        doc,
        "使用方式",
        "課堂中先以鉛筆完成填空與圖像判讀；每一節末再用「隨堂檢核」確認自己能否說出原因，而不只背結論。",
    )

    add_section_heading(doc, "本章學習目標", 2)
    bullet_id = add_numbering_definition(doc, ordered=False)
    goals = [
        "能用右手定則判斷直導線、圓形線圈與螺線管的磁場方向。",
        "能用冷次定律判斷感應磁場、感應電流與磁鐵受力方向。",
        "能說明電磁波的來源、性質與光譜排列。",
        "能運用 v = fλ、反射／折射、繞射與干涉的核心觀念。",
        "能以相對運動判斷都卜勒效應中的頻率升降。",
    ]
    for goal in goals:
        add_list_item(doc, "□ " + goal, bullet_id)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("資料來源：教師 Canva〈高一物理：CH.4〉課程內容")
    set_run_font(r, size=8.5, color=MUTED, italic=True)


def section_current_magnetism(doc: Document, images):
    doc.add_page_break()
    add_section_heading(doc, "一、電流磁效應：動電生磁", 1)
    add_callout(doc, "核心問題", "電流如何產生磁場？我們如何用手勢把三維方向讀出來？")

    add_section_heading(doc, "1. 厄斯特的發現", 2)
    add_fill(
        doc,
        "厄斯特發現：放在載流導線旁的磁針會發生 __________，顯示「__________ 能產生 __________」。",
        1,
    )
    add_body(
        doc,
        "約定：物理題中的「電流方向」指正電荷移動方向，因此與電子流方向 __________。",
        after=7,
    )

    add_section_heading(doc, "2. 直導線的安培右手定則", 2)
    num_id = add_numbering_definition(doc, ordered=True)
    for text in [
        "右手大拇指指向 __________ 的方向。",
        "彎曲的四指指出導線周圍 __________ 的環繞方向。",
        "若從某端看見磁力線逆時針旋轉，該端電流方向應為 __________（朝向／背向觀察者）。",
    ]:
        add_list_item(doc, text, num_id)
    add_picture(doc, images["wire"], width=6.35)
    add_caption(doc, "圖 1　請先判讀 ⊙、⊗ 的意義，再完成兩側磁場方向。")

    add_table(
        doc,
        ["空間記號", "代表方向", "生活化記憶"],
        [
            ["⊙", "________________", "像箭頭的 __________ 朝向你"],
            ["⊗", "________________", "像箭尾的 __________ 背向你"],
        ],
        [1600, 2900, 4860],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_section_heading(doc, "3. 圓形線圈、螺線管與電磁鐵", 2)
    add_fill(
        doc,
        "載流圓形線圈手勢：四指沿著 __________ 方向彎曲；大拇指指出線圈中心的 __________ 方向，也指向線圈的 __________ 極。",
        1,
    )
    add_fill(
        doc,
        "載流螺線管可視為很多圈的 __________。電流越大、單位長度匝數越 __________，內部磁場通常越強。",
        1,
    )
    add_table(
        doc,
        ["類型", "磁性特徵", "常見例子／用途"],
        [
            ["硬磁鐵", "本身可長期保有磁性，屬於 __________ 磁鐵", "指南針、永磁馬達"],
            ["軟磁鐵", "會配合外加磁場磁化；移除磁場後較易失去磁性", "鐵、鈷、鎳；電磁鐵的 __________"],
        ],
        [1700, 4250, 3410],
    )
    add_callout(
        doc,
        "電磁鐵",
        "螺線管中加入軟磁性鐵心，可增強磁場；最大優點是能用「通電／斷電」控制磁力。",
        fill="FFF8E8",
        color=ACCENT,
    )

    add_section_heading(doc, "4. 載流直導線在磁場中的受力", 2)
    add_fill(
        doc,
        "當電流方向與磁場方向互相 __________ 時，磁力最明顯；若兩者互相 __________，則磁力為 0。",
        1,
    )
    add_body(
        doc,
        "定性關係：磁力大小會隨電流 I、磁場 B 增大而 __________。方向可用課堂中的右手開掌法判斷。",
        after=7,
    )

    add_question_group(
        doc,
        "隨堂檢核 A",
        [
            ("一條直導線中的傳統電流垂直紙面向外。請畫出導線周圍的磁場方向，並寫出判斷手勢。", 2),
            ("從某圓形線圈正面看，電流為逆時針。此面是 N 極還是 S 極？線圈中心磁場朝哪裡？", 1),
            ("請列出三種能使螺線管磁場增強的方法。", 2),
        ],
    )


def section_induction(doc: Document, images):
    doc.add_page_break()
    add_section_heading(doc, "二、電磁感應：動磁生電", 1)
    add_callout(doc, "核心問題", "只要有磁場就會生電嗎？真正關鍵是「穿過線圈的磁通量是否改變」。")

    add_section_heading(doc, "1. 法拉第與磁通量", 2)
    add_fill(
        doc,
        "法拉第發現：磁鐵相對線圈靠近或遠離時，檢流計會偏轉；因電磁感應而產生的電流稱為 __________。",
        1,
    )
    add_body(doc, "本課以磁場垂直線圈平面時為主：　Φ = B × S", bold=True, color=DARK_BLUE)
    add_table(
        doc,
        ["符號", "物理量", "意義／單位"],
        [
            ["Φ", "磁通量", "線圈「吃到」的磁力線數量；單位 __________"],
            ["B", "磁場", "磁場強弱；單位 __________"],
            ["S", "面積", "線圈有效面積；單位 __________"],
        ],
        [1200, 2100, 6060],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_callout(
        doc,
        "判斷先後",
        "先判斷原磁通量如何改變，再決定線圈要產生哪個方向的感應磁場；最後才用線圈右手定則找電流。",
    )

    add_section_heading(doc, "2. 冷次定律（Lenz's law）", 2)
    add_fill(
        doc,
        "冷次定律：感應電流產生的磁場，恆會 __________ 磁通量的「變化」。它反抗的是變化，不一定反抗原磁場本身。",
        1,
    )
    num_id = add_numbering_definition(doc, ordered=True)
    steps = [
        "辨認原磁場穿過線圈的方向：__________。",
        "判斷磁通量變化：增加／減少／不變。",
        "線圈產生感應磁場，目的是 __________ 這個變化。",
        "用載流圓形線圈右手定則，決定感應電流方向。",
    ]
    for step in steps:
        add_list_item(doc, step, num_id)
    add_picture(doc, images["lenz"], width=6.35)
    add_caption(doc, "圖 2　磁鐵只有平移；請完成線圈靠近磁鐵一側的極性與電流方向。")

    add_table(
        doc,
        ["情況（N 極面向線圈）", "磁通量", "線圈靠磁鐵側", "由磁鐵側看電流"],
        [
            ["N 極靠近", "增加", "__________ 極", "__________ 時針"],
            ["N 極遠離", "減少", "__________ 極", "__________ 時針"],
            ["磁鐵與線圈相對靜止", "__________", "無感應磁場", "無感應電流"],
        ],
        [2800, 1800, 2380, 2380],
        font_size=9.2,
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )

    add_section_heading(doc, "3. 電磁感應的應用", 2)
    add_table(
        doc,
        ["裝置", "造成磁通量改變的方法", "能量／功能"],
        [
            ["交流發電機", "線圈在磁場中 __________", "__________ 能 → 電能"],
            ["電磁爐", "交流電造成 __________ 磁場", "鍋底感應電流的 __________ 效應使鍋子發熱"],
            ["無線充電器", "發射線圈建立時變磁場", "接收線圈產生 __________，再替電池充電"],
        ],
        [1900, 3300, 4160],
    )

    add_question_group(
        doc,
        "隨堂檢核 B",
        [
            ("一個線圈完全位於均勻且恆定的磁場內，以固定速度平移，但線圈面積與方向都不變。是否產生感應電流？請說明。", 2),
            ("若把磁鐵的 S 極朝線圈靠近，線圈靠磁鐵的一側應成為哪一極？由磁鐵側看，電流是順時針或逆時針？", 2),
            ("為什麼冷次定律的結果常使移動中的磁鐵減速？請用能量觀點簡述。", 2),
        ],
    )


def section_em_waves(doc: Document):
    doc.add_page_break()
    add_section_heading(doc, "三、電與磁的統整：電磁波", 1)
    add_callout(doc, "核心問題", "若變動的電場與磁場能彼此產生，它們能否離開電荷，在空間中自行傳播？")

    add_section_heading(doc, "1. 馬克士威的預測與赫茲的證實", 2)
    add_fill(
        doc,
        "馬克士威統整前人的實驗結果，建立四大電磁方程式，預測電場與磁場可彼此 __________ 並向外傳播。",
        1,
    )
    add_fill(
        doc,
        "計算得到此波在真空中的速度與 __________ 相同，因此提出「光也是一種 __________」；後來由 __________ 實驗證實。",
        1,
    )
    add_body(doc, "產生電磁波的關鍵：電荷必須做 __________ 運動，才能形成持續變動的電場。")

    add_section_heading(doc, "2. 電磁波的共同性質", 2)
    bullet_id = add_numbering_definition(doc, ordered=False)
    properties = [
        "傳播不需要 __________，因此可在真空中前進。",
        "電場振盪方向、磁場振盪方向、波的前進方向彼此 __________。",
        "屬於 __________ 波。",
        "在真空中各種電磁波的波速相同，但波長與頻率不同。",
    ]
    for item in properties:
        add_list_item(doc, item, bullet_id)

    add_section_heading(doc, "3. 電磁波光譜", 2)
    add_callout(
        doc,
        "排序線",
        "波長由大到小：無線電波 → 微波 → 紅外線 → 可見光 → 紫外線 → X 射線 → γ 射線",
        fill="EEF8F4",
        color=GREEN,
    )
    add_fill(
        doc,
        "沿上列由左向右，頻率逐漸 __________，單一光子的能量逐漸 __________，繞射能力通常逐漸 __________。",
        1,
    )
    add_table(
        doc,
        ["光譜區域", "典型應用或現象", "主要特徵"],
        [
            ["無線電波／微波", "通訊、雷達、加熱", "波長較 __________，較容易繞射"],
            ["紅外線／可見光", "熱像、遙控、視覺", "可見光只是電磁光譜中的一小段"],
            ["紫外線／X 射線／γ 射線", "殺菌、醫療影像、治療", "頻率較 __________、能量較高"],
        ],
        [2300, 3400, 3660],
    )

    add_section_heading(doc, "4. 光的微粒說與波動說", 2)
    add_table(
        doc,
        ["年代／人物", "主張或實驗", "請填關鍵詞"],
        [
            ["1672／1704　牛頓", "提出光的微粒說；以三稜鏡研究白光", "白光由 __________ 組成"],
            ["1678　惠更斯", "提出光的波動說", "光可發生繞射與 __________"],
            ["1803　楊格", "雙狹縫實驗", "觀察到光的 __________，支持波動說"],
            ["1849／1850　菲左、傅科", "測量光速；比較水中與真空光速", "水中光速較 __________"],
            ["1865　馬克士威", "由電磁理論推導", "光是一種 __________"],
            ["1905　愛因斯坦", "以光子解釋光電效應", "光具有波粒 __________"],
        ],
        [2550, 4010, 2800],
        font_size=8.9,
    )

    add_question_group(
        doc,
        "隨堂檢核 C",
        [
            ("請由頻率低到高寫出完整的電磁波光譜。", 2),
            ("為什麼加速電荷能產生電磁波，而靜止電荷不能？", 2),
            ("「光到底是粒子還是波？」請用兩個實驗證據回答。", 3),
        ],
    )


def section_waves_optics(doc: Document, images):
    doc.add_page_break()
    add_section_heading(doc, "四、波動與光學", 1)
    add_callout(doc, "核心問題", "波只傳能量，介質粒子只在平衡位置附近振動；所有光學現象都可回到波的傳播方式理解。")

    add_section_heading(doc, "1. 波的基本物理量", 2)
    add_table(
        doc,
        ["物理量", "符號／單位", "定義（請填空）"],
        [
            ["週期", "T／s", "振動 __________ 次所需的時間"],
            ["頻率", "f／Hz", "每 __________ 振動的次數"],
            ["波長", "λ／m", "相鄰兩個同相位點的 __________"],
            ["振幅", "A／m", "波峰或波谷到 __________ 位置的距離"],
            ["波速", "v／m·s⁻¹", "波形／能量在介質中傳遞的 __________"],
        ],
        [1500, 1900, 5960],
        font_size=9.3,
    )
    add_body(doc, "公式整理：　T = 1 / f　　v = f λ", bold=True, color=DARK_BLUE, size=12)
    add_fill(
        doc,
        "波進入另一種介質時，頻率由 __________ 決定，因此通常保持不變；改變的是波速與 __________。",
        1,
    )
    add_picture(doc, images["wave"], width=6.35)
    add_caption(doc, "圖 3　請在圖中標出波長與振幅，並用箭頭畫出一個質點此刻可能的振動方向。")
    add_body(
        doc,
        "強度與振幅的關係：波的強度 ∝ A²。因此振幅變為 2 倍時，強度變為 __________ 倍。",
    )

    add_section_heading(doc, "2. 橫波與縱波", 2)
    add_table(
        doc,
        ["波形", "介質振動方向與波前進方向", "例子"],
        [
            ["橫波", "互相 __________", "繩波、電磁波"],
            ["縱波", "互相 __________", "空氣中的聲波"],
        ],
        [1800, 4300, 3260],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_section_heading(doc, "3. 光的反射與折射", 2)
    add_fill(
        doc,
        "反射定律：入射角 = __________；角度都以 __________ 為基準測量。光路具有 __________ 性。",
        1,
    )
    add_body(doc, "折射率：n = c / v　　司乃爾定律：n₁ sin θ₁ = n₂ sin θ₂",
             bold=True, color=DARK_BLUE, size=11.5)
    add_fill(
        doc,
        "光由快介質進入慢介質時會偏向 __________；由慢介質進入快介質時會 __________ 法線。折射率越大，介質中的光速越 __________。",
        1,
    )
    add_picture(doc, images["optics"], width=6.35)
    add_caption(doc, "圖 4　左：反射；右：折射。請補上角度與介質折射率。")
    add_callout(
        doc,
        "生活連結",
        "海市蜃樓來自不同高度空氣溫度不同，造成折射率逐層改變，光線因此連續彎曲。",
        fill="FFF8E8",
        color=ACCENT,
    )

    add_section_heading(doc, "4. 波前與惠更斯原理", 2)
    add_fill(
        doc,
        "波前是所有同相位點的連線，與波的行進方向互相 __________。實線常代表 __________ 連線，虛線可代表波谷連線。",
        1,
    )
    add_fill(
        doc,
        "惠更斯原理：波前上的每一點都可視為 __________，各自發出子波；所有子波的 __________ 線形成下一時刻的波前。",
        1,
    )

    add_section_heading(doc, "5. 繞射與干涉", 2)
    add_fill(
        doc,
        "當波長 λ 與狹縫寬度或障礙物尺寸 __________ 時，繞射最明顯。可見光波長很小，因此日常看起來近似 __________ 傳播。",
        1,
    )
    add_fill(
        doc,
        "兩個以上的波相遇時，位移會彼此 __________，形成合成波。完全建設性干涉時振幅 __________；完全破壞性干涉時位移互相 __________。",
        1,
    )
    add_body(
        doc,
        "光的干涉需使用同調光：頻率相同且相位差保持穩定。薄膜干涉則來自光在不同介面反射後彼此干涉。",
        after=7,
    )
    add_picture(doc, images["diffraction_doppler"], width=6.35)
    add_caption(doc, "圖 5　左圖完成繞射結論；右圖將在下一節用來判讀都卜勒效應。")

    add_question_group(
        doc,
        "隨堂檢核 D",
        [
            ("一列波的頻率為 5.0 Hz，波長為 0.80 m。求波速。", 2),
            ("光由空氣斜射入玻璃。請畫出折射光線，並比較 θ空氣 與 θ玻璃 的大小。", 2),
            ("為什麼隔著牆容易聽見低沉的聲音，卻不容易『繞射看見』牆後物體？", 2),
            ("兩個振幅相同、方向相反的脈波相遇時會暫時消失。相遇後它們會不會永久消失？請說明。", 2),
        ],
    )


def section_doppler(doc: Document):
    doc.add_page_break()
    add_section_heading(doc, "五、都卜勒效應", 1)
    add_callout(doc, "一行統整", "不論誰在移動：相對靠近 → 接收頻率升高；相對遠離 → 接收頻率降低；相對靜止 → 頻率不變。")

    add_section_heading(doc, "1. 波源動、觀察者動，有什麼不同？", 2)
    add_body(
        doc,
        "理想情況下，聲源本身發出的頻率不變；都卜勒效應是觀察者接收到波前的速率改變。",
    )
    add_table(
        doc,
        ["移動者", "直接改變的量", "為什麼"],
        [
            ["波源", "波前間距，也就是 __________", "波源在相鄰兩次發波之間已移動"],
            ["觀察者", "與波前的相對速度／接收 __________", "觀察者迎向或遠離既有波前"],
        ],
        [1800, 3000, 4560],
    )
    add_fill(
        doc,
        "聲速由 __________ 決定；同一環境中，即使聲源移動，聲波相對介質的速率仍不變。",
        1,
    )

    add_section_heading(doc, "2. 四種基本情境", 2)
    add_table(
        doc,
        ["情境", "接收頻率", "波源造成的波長", "觀察者感受的相對波速"],
        [
            ["波源靠近觀察者", "__________", "變 __________", "不因波源而改變"],
            ["波源遠離觀察者", "__________", "變 __________", "不因波源而改變"],
            ["觀察者靠近波源", "__________", "不變", "感覺變 __________"],
            ["觀察者遠離波源", "__________", "不變", "感覺變 __________"],
        ],
        [2460, 1800, 2500, 2600],
        font_size=8.9,
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    add_callout(
        doc,
        "容易混淆",
        "『感覺波速改變』指波相對觀察者的接近速度改變；聲波相對空氣的真正波速仍由介質決定。",
        fill="FFF8E8",
        color=ACCENT,
    )

    add_section_heading(doc, "3. 兩者都動的快速判讀", 2)
    add_table(
        doc,
        ["相對狀態", "接收頻率", "判讀"],
        [
            ["彼此靠近", "更 __________", "波長壓縮／迎向波前，兩效果同向"],
            ["彼此遠離", "更 __________", "波長拉長／遠離波前，兩效果同向"],
            ["同方向、同速率", "__________", "彼此相對靜止；兩效果互相抵銷"],
        ],
        [2400, 2100, 4860],
    )

    add_section_heading(doc, "4. 光的都卜勒效應", 2)
    add_fill(
        doc,
        "天體相對地球遠離時，接收到的光頻率降低、波長變長，稱為 __________；相對靠近時則稱為 __________。",
        1,
    )
    add_body(
        doc,
        "氣象雷達與天文光譜都能利用頻率偏移，反推出目標沿視線方向的運動。",
    )

    add_section_heading(doc, "5. 相對速度補充（選修）", 2)
    add_body(doc, "A 相對於 B 的速度：　v_AB = v_A − v_B", bold=True, color=DARK_BLUE, size=11.5)
    add_fill(
        doc,
        "若向右為正，A 以 −10 m/s、B 以 +10 m/s 運動，則 v_BA = __________，表示 A 看見 B 向 __________ 以 __________ m/s 移動。",
        2,
    )

    add_question_group(
        doc,
        "隨堂檢核 E",
        [
            ("救護車鳴笛並朝你駛來。比較你聽到的頻率與原頻率，並說明波長如何改變。", 2),
            ("你騎車遠離固定蜂鳴器。聲源前方波長是否改變？你接收到的頻率為何下降？", 2),
            ("波源在後、觀察者在前，兩者同方向且以相同速率前進。接收頻率會改變嗎？請分別從波長與相對波速說明。", 3),
            ("某星系的吸收光譜相對實驗室光譜向紅端位移。該星系相對地球如何運動？", 1),
        ],
    )


def section_synthesis(doc: Document):
    doc.add_page_break()
    add_section_heading(doc, "六、章末統整與自我檢核", 1)
    add_callout(doc, "把本章串起來", "電流產生磁場；改變的磁通量產生感應電流；變動的電場與磁場能以電磁波傳播，而光正是其中一員。")

    add_section_heading(doc, "1. 概念鏈填空", 2)
    chain = doc.add_table(rows=3, cols=5)
    widths = [1872] * 5
    set_table_geometry(chain, widths, indent_dxa=0)
    set_table_borders(chain, color="C8D5E4", size=6)
    top = ["電流", "磁場", "磁通量變化", "感應電流", "電磁波／光"]
    bottom = [
        "透過 ________ 定則",
        "可作用於載流導線",
        "由 ________ 定律判方向",
        "可用於發電與加熱",
        "具有波動與 ________ 性",
    ]
    for i in range(5):
        set_cell_text(chain.rows[0].cells[i], top[i], bold=True, color=WHITE,
                      size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER, fill=BLUE)
        set_cell_text(chain.rows[1].cells[i], "→" if i < 4 else "◎",
                      bold=True, color=ACCENT, size=15,
                      align=WD_ALIGN_PARAGRAPH.CENTER, fill=WHITE)
        set_cell_text(chain.rows[2].cells[i], bottom[i], size=8.6,
                      align=WD_ALIGN_PARAGRAPH.CENTER, fill=LIGHTER_BLUE)
    set_table_geometry(chain, widths, indent_dxa=0)
    add_body(doc, "", after=2)

    add_section_heading(doc, "2. 綜合練習", 2)
    questions = [
        ("（選擇）直導線電流垂直紙面向內，紙面上的磁場方向為：（A）順時針（B）逆時針（C）向上（D）向下", 1),
        ("（選擇）要增強螺線管磁場，下列何者無效？（A）增大電流（B）增加匝數密度（C）加入軟鐵心（D）使電流變為零", 1),
        ("（簡答）電子向右運動時，傳統電流方向為何？若要使用安培右手定則，拇指應指哪裡？", 2),
        ("（選擇）磁鐵在固定線圈旁靜止不動時：（A）磁通量大所以有感應電流（B）磁通量不變所以沒有感應電流（C）一定有交流電（D）無法判斷", 1),
        ("（推理）N 極離開線圈。請依序寫出：磁通量變化 → 感應磁場方向 → 線圈近端極性 → 電流方向。", 3),
        ("（選擇）下列何者不是電磁波？（A）可見光（B）X 射線（C）聲波（D）微波", 1),
        ("（排序）將紅外線、無線電波、紫外線、可見光依頻率由低到高排列。", 2),
        ("（計算）波速 12 m/s、頻率 3.0 Hz，求波長與週期。", 2),
        ("（判斷）波從細繩進入粗繩後頻率變為原來一半。正確／錯誤？請說明。", 2),
        ("（作圖）畫出光由玻璃斜射入空氣的折射方向，並標示法線、入射角與折射角。", 3),
        ("（選擇）繞射最明顯的條件是：（A）λ 遠小於狹縫（B）λ 約等於狹縫（C）頻率最高（D）振幅為零", 1),
        ("（簡答）雙狹縫為什麼能讓兩束光具有同調性？", 2),
        ("（都卜勒）固定觀察者聽見遠離中的聲源，其音高如何？波長如何？聲速如何？", 2),
        ("（都卜勒）固定聲源不變，觀察者迎向聲源。哪個量直接改變：波長、相對波速、接收頻率？", 2),
        ("（跨章）用一段話說明「發電機如何把力學能轉成電能」，至少使用：磁通量、感應電流、冷次定律。", 4),
    ]
    num_id = add_numbering_definition(doc, ordered=True)
    for text, lines in questions:
        p = add_list_item(doc, text, num_id)
        p.paragraph_format.space_after = Pt(3)
        for _ in range(lines):
            ans = doc.add_paragraph()
            ans.paragraph_format.left_indent = Inches(0.38)
            ans.paragraph_format.right_indent = Inches(0.1)
            ans.paragraph_format.space_before = Pt(3)
            ans.paragraph_format.space_after = Pt(5)
            ans.add_run("\u00A0")
            set_paragraph_border(ans, "bottom", color="B8C2CF", size=5, space=2)

    add_section_heading(doc, "3. 我的學習狀態", 2)
    add_table(
        doc,
        ["我能做到…", "很穩", "需複習", "想提問"],
        [
            ["用手勢判斷電流與磁場方向", "□", "□", "□"],
            ["用冷次定律說出完整推理鏈", "□", "□", "□"],
            ["排列電磁波光譜並比較 λ、f、能量", "□", "□", "□"],
            ["運用 v = fλ、反射／折射與波動觀念", "□", "□", "□"],
            ["以相對靠近／遠離判斷都卜勒效應", "□", "□", "□"],
        ],
        [5660, 1233, 1233, 1234],
        font_size=9.3,
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    add_fill(doc, "我最需要再問老師的一個問題：", 3)
    add_fill(doc, "本章我最有把握的一個概念，以及我的理由：", 3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("— 完成不是終點；能解釋「為什麼」才是真的理解。—")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)


def set_core_properties(doc: Document):
    props = doc.core_properties
    props.title = "高一物理 CH.4 電與磁的統一－學生學習單"
    props.subject = "電流磁效應、電磁感應、電磁波、波動與光學、都卜勒效應"
    props.author = "洪瑋澤老師課程內容整理"
    props.keywords = "高一物理, 電與磁, 學習單, 學生講義"
    props.comments = "依據 Canva〈高一物理：CH.4〉內容製作之學生用填寫版學習單。"


def audit_tables(doc: Document):
    for idx, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa":
            raise RuntimeError(f"Table {idx} lacks fixed DXA tblW")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise RuntimeError(f"Table {idx} lacks fixed layout")
        if any(cell._tc.get_or_add_tcPr().find(qn("w:tcW")) is None
               for row in table.rows for cell in row.cells):
            raise RuntimeError(f"Table {idx} has a cell without tcW")


def main():
    images = build_diagrams()
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    section_current_magnetism(doc, images)
    section_induction(doc, images)
    section_em_waves(doc)
    section_waves_optics(doc, images)
    section_doppler(doc)
    section_synthesis(doc)
    audit_tables(doc)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
